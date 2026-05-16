from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_DOCX = DOCS / "AtomQuest_Hackathon_Submission_Packet.docx"
OUT_PDF = DOCS / "AtomQuest_Hackathon_Submission_Packet.pdf"
DIAGRAM_PNG = DOCS / "architecture_diagram.png"

LIVE_URL = "https://atomquest-907l.onrender.com/"
SOURCE_URL = "https://github.com/coder200412/Atomquet"


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw, xy, title, lines, fill="#ffffff", outline="#d7dee9"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    draw.text((x1 + 24, y1 + 22), title, fill="#111827", font=font(22, True))
    y = y1 + 58
    for line in lines:
        draw.text((x1 + 24, y), line, fill="#344054", font=font(16))
        y += 26


def arrow(draw, start, end):
    draw.line([start, end], fill="#64748b", width=3)
    sx, sy = start
    ex, ey = end
    if ex > sx:
        points = [(ex, ey), (ex - 12, ey - 7), (ex - 12, ey + 7)]
    elif ex < sx:
        points = [(ex, ey), (ex + 12, ey - 7), (ex + 12, ey + 7)]
    elif ey > sy:
        points = [(ex, ey), (ex - 7, ey - 12), (ex + 7, ey - 12)]
    else:
        points = [(ex, ey), (ex - 7, ey + 12), (ex + 7, ey + 12)]
    draw.polygon(points, fill="#64748b")


def create_diagram():
    img = Image.new("RGB", (1400, 920), "#f5f7fb")
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "AtomQuest Goal Setting and Tracking Portal", fill="#111827", font=font(32, True))
    draw.text((70, 88), "React + Express + Prisma + Render Postgres, hosted on Render", fill="#667085", font=font(18))

    draw_box(draw, (70, 150, 330, 280), "Web Browser", ["Public demo URL"])
    draw_box(draw, (390, 150, 680, 280), "Render Web Service", ["Node + Express", "Serves React build"], fill="#edf9f5", outline="#0f766e")
    draw_box(draw, (740, 150, 1020, 280), "Render Postgres", ["Managed database", "Private connection"])

    arrow(draw, (330, 215), (390, 215))
    arrow(draw, (680, 215), (740, 215))

    draw_box(draw, (70, 390, 340, 535), "Employee", ["Draft goals", "Submit goals", "Quarterly check-ins"])
    draw_box(draw, (390, 390, 680, 535), "Manager L1", ["Review and edit", "Approve or return", "Check-in comments"])
    draw_box(draw, (740, 390, 1020, 535), "Admin / HR", ["Cycle controls", "Unlock exceptions", "Audit and exports"])

    arrow(draw, (535, 280), (535, 345))
    draw.line([(535, 345), (205, 345), (205, 390)], fill="#64748b", width=3)
    draw.polygon([(205, 390), (198, 378), (212, 378)], fill="#64748b")
    arrow(draw, (535, 345), (535, 390))
    draw.line([(535, 345), (880, 345), (880, 390)], fill="#64748b", width=3)
    draw.polygon([(880, 390), (873, 378), (887, 378)], fill="#64748b")

    draw_box(draw, (70, 660, 340, 800), "Microsoft Entra", ["SSO simulation", "Group-role mapping"])
    draw_box(draw, (390, 660, 680, 800), "Email + Teams", ["Notifications", "Adaptive cards"])
    draw_box(draw, (740, 660, 1020, 800), "Escalations", ["Rule engine", "Admin logs"])
    draw_box(draw, (1070, 660, 1340, 800), "Analytics", ["QoQ trends", "Completion heatmaps"])

    arrow(draw, (535, 535), (535, 610))
    draw.line([(535, 610), (205, 610), (205, 660)], fill="#64748b", width=3)
    draw.polygon([(205, 660), (198, 648), (212, 648)], fill="#64748b")
    arrow(draw, (535, 610), (535, 660))
    draw.line([(535, 610), (880, 610), (880, 660)], fill="#64748b", width=3)
    draw.polygon([(880, 660), (873, 648), (887, 648)], fill="#64748b")
    draw.line([(535, 610), (1205, 610), (1205, 660)], fill="#64748b", width=3)
    draw.polygon([(1205, 660), (1198, 648), (1212, 648)], fill="#64748b")

    img.save(DIAGRAM_PNG, "PNG")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1"
    paragraph.add_run(text)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "EAF3F1")
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def create_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.color.rgb = RGBColor(15, 118, 110)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AtomQuest Goal Portal - Hackathon Submission")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("In-House Goal Setting & Tracking Portal")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 112, 133)

    add_heading(doc, "Submission Deliverables")
    add_table(
        doc,
        ["#", "Deliverable", "Details"],
        [
            ["1", "Live / hosted demo URL", LIVE_URL],
            ["2", "Source code repository", SOURCE_URL],
            ["3", "Architecture diagram", "Included below and available at docs/architecture_diagram.png"],
            ["4", "Role credentials / user journeys", "Demo credentials below; public Employee/Manager signup is enabled."],
        ],
        [0.35, 2.0, 4.4],
    )

    add_heading(doc, "Login Credentials")
    para = doc.add_paragraph("All seeded demo accounts use password123. Visitors can also create a new Employee or Manager account from the login screen.")
    para.paragraph_format.space_after = Pt(6)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ["Admin / HR", "admin@atomquest.local", "password123"],
            ["Manager", "maya.manager@atomquest.local", "password123"],
            ["Manager", "karan.manager@atomquest.local", "password123"],
            ["Employee", "neha.employee@atomquest.local", "password123"],
            ["Employee", "arjun.employee@atomquest.local", "password123"],
            ["Employee", "diya.employee@atomquest.local", "password123"],
        ],
        [1.35, 3.3, 1.35],
    )

    add_heading(doc, "Architecture Diagram")
    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(DIAGRAM_PNG), width=Inches(7.0))

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "BRD Coverage Summary")
    add_table(
        doc,
        ["Area", "Implemented Coverage"],
        [
            ["Goal creation and validation", "Goal title, description, thrust area, UoM, target, and weightage with 10% minimum, 8-goal maximum, and 100% total checks."],
            ["Approval workflow", "Managers review submitted sheets, edit during review, approve, or return for rework. Approved sheets are locked until Admin / HR unlock."],
            ["Quarterly check-ins", "Employees log actual achievements and status by quarter; managers add structured comments."],
            ["Reporting and governance", "Admin dashboard, CSV export, completion tracking, audit trail, cycle override, and exception unlock."],
            ["Bonus modules", "Microsoft Entra simulation, Email/Teams notifications, escalation rules, adaptive cards, QoQ trends, heatmaps, and manager effectiveness."],
        ],
        [2.2, 4.8],
    )

    add_heading(doc, "Render Deployment")
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Build Command", "npm install && npx prisma generate && npm run build"],
            ["Start Command", "npm run db:deploy && npm run start"],
            ["Database", "Render Postgres using DATABASE_URL internal connection string"],
            ["Public Signup", "Enabled for Employee and Manager accounts"],
        ],
        [1.8, 5.2],
    )

    doc.save(OUT_DOCX)


def pdf_table(headers, rows, widths):
    data = [headers] + rows
    table = Table(data, colWidths=[width * inch for width in widths], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF3F1")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CFD7E3")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def create_pdf():
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "AQTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        textColor=colors.HexColor("#111827"),
        alignment=1,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "AQSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        textColor=colors.HexColor("#667085"),
        alignment=1,
        spaceAfter=16,
    )
    heading = ParagraphStyle(
        "AQHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=colors.HexColor("#0F766E"),
        spaceBefore=10,
        spaceAfter=8,
    )
    body = ParagraphStyle(
        "AQBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#28313D"),
        spaceAfter=7,
    )

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = [
        Paragraph("AtomQuest Goal Portal - Hackathon Submission", title),
        Paragraph("In-House Goal Setting & Tracking Portal", subtitle),
        Paragraph("Submission Deliverables", heading),
        pdf_table(
            ["#", "Deliverable", "Details"],
            [
                ["1", "Live / hosted demo URL", LIVE_URL],
                ["2", "Source code repository", SOURCE_URL],
                ["3", "Architecture diagram", "Included below and available at docs/architecture_diagram.png"],
                ["4", "Role credentials / user journeys", "Demo credentials below; public Employee/Manager signup is enabled."],
            ],
            [0.35, 1.85, 5.0],
        ),
        Spacer(1, 10),
        Paragraph("Login Credentials", heading),
        Paragraph("All seeded demo accounts use <b>password123</b>. Visitors can also create a new Employee or Manager account from the login screen.", body),
        pdf_table(
            ["Role", "Email", "Password"],
            [
                ["Admin / HR", "admin@atomquest.local", "password123"],
                ["Manager", "maya.manager@atomquest.local", "password123"],
                ["Manager", "karan.manager@atomquest.local", "password123"],
                ["Employee", "neha.employee@atomquest.local", "password123"],
                ["Employee", "arjun.employee@atomquest.local", "password123"],
                ["Employee", "diya.employee@atomquest.local", "password123"],
            ],
            [1.3, 3.5, 1.2],
        ),
        Spacer(1, 10),
        Paragraph("Architecture Diagram", heading),
        PdfImage(str(DIAGRAM_PNG), width=7.2 * inch, height=4.42 * inch),
        PageBreak(),
        Paragraph("BRD Coverage Summary", heading),
        pdf_table(
            ["Area", "Implemented Coverage"],
            [
                ["Goal creation and validation", "Goal title, description, thrust area, UoM, target, and weightage with 10% minimum, 8-goal maximum, and 100% total checks."],
                ["Approval workflow", "Managers review submitted sheets, edit during review, approve, or return for rework. Approved sheets are locked until Admin / HR unlock."],
                ["Quarterly check-ins", "Employees log actual achievements and status by quarter; managers add structured comments."],
                ["Reporting and governance", "Admin dashboard, CSV export, completion tracking, audit trail, cycle override, and exception unlock."],
                ["Bonus modules", "Microsoft Entra simulation, Email/Teams notifications, escalation rules, adaptive cards, QoQ trends, heatmaps, and manager effectiveness."],
            ],
            [1.9, 5.2],
        ),
        Spacer(1, 10),
        Paragraph("Render Deployment", heading),
        pdf_table(
            ["Setting", "Value"],
            [
                ["Build Command", "npm install && npx prisma generate && npm run build"],
                ["Start Command", "npm run db:deploy && npm run start"],
                ["Database", "Render Postgres using DATABASE_URL internal connection string"],
                ["Public Signup", "Enabled for Employee and Manager accounts"],
            ],
            [1.55, 5.55],
        ),
    ]
    doc.build(story)


def main():
    DOCS.mkdir(exist_ok=True)
    create_diagram()
    create_docx()
    create_pdf()
    print(OUT_DOCX)
    print(OUT_PDF)


if __name__ == "__main__":
    main()
